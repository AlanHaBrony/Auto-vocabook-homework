'''
This demo now only fit the homework of my school, but its method can be generalize
And it is only used to avoid boring ctlC and ctlV

!You should make sure you truly remember the words!
Good day.

Alex Zhang, 22 Oct. 2021

Emm it seems to be a little bit time-consuming. Anyway,still save a lot.
And be sure to type the words correctly._(:3l/_)_
(22.03-)

'''

import docx
from docx import Document
import requests
import re
from bs4 import BeautifulSoup
from urllib import request

class autowork(object):
    def __init__(self):
        self.sufurla='http://dict.youdao.com/search?keyfrom=webwordbook&q='
        self.sufurlb='http://wordnetweb.princeton.edu/perl/webwn?s='
        self.file=docx.Document("Alex Zhang's vocab.docx")
    
    def get_phonetic(self,word):
        r0=requests.get(self.sufurla+'word')
        html0=r0.text
        soup=BeautifulSoup(html0,'html.parser')
        s=soup.find('span',attrs={'class':'phonetic'})
        phonetic=re.findall(r"(\[.*?\])",s.text)
        return(phonetic[0])
    
    def get_meaning(self,word):
        r1=requests.get(self.sufurlb+'word')
        html1=r1.text
        soup1=BeautifulSoup(html1,'html.parser')
        prim1=soup1.find('li').get_text
        meaning=re.findall('[(](.*?)[)]',str(prim1))
        return(meaning[1])
    
    def get_example(self,word):
        r2=requests.get(self.sufurlb+'word')
        html2=r2.text
        soup2=BeautifulSoup(html2,'html.parser')
        prim2=soup2.find('i').get_text()
        separate=re.findall('\w+',prim2)
        example=' '.join(separate)+"."
        return(example)
    
    def autowrite(self):
        table=self.file.tables
        for i in range(26):
            word=table[1].rows[i].cells[0].text
            table[1].rows[i].cells[1].text=self.get_phonetic(word)
            table[1].rows[i].cells[2].text=self.get_meaning(word)
            table[1].rows[i].cells[3].text=self.get_example(word)
        self.file.save("Alex Zhang's vocab.docx")

        
    def clear(self):
        #default False
        table=self.file.tables
        for i in range(25):
            for j in range(5):
                table[1].rows[i].cells[j].text=' '
        self.file.save("Alex Zhang's vocab.docx")

if __name__=='__main__':
    ans=autowork()
    
    ans.autowrite()


            
        
